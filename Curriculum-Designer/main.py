import sys
import os
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from curriculum_designer import CurriculumDesignerSystem

def save_to_json(data, topic, target_audience, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_topic = "".join(c for c in topic if c.isalnum()).lower()
    safe_audience = "".join(c for c in target_audience if c.isalnum()).lower()
    file_prefix = f"{safe_topic}_{safe_audience}"
    json_path = os.path.join(output_dir, f"{file_prefix}_{timestamp}.json")
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    print(f"🗄️  Saved JSON to: {json_path}")

def save_to_docx(module, topic, target_audience, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_topic = "".join(c for c in topic if c.isalnum()).lower()
    safe_audience = "".join(c for c in target_audience if c.isalnum()).lower()
    file_prefix = f"{safe_topic}_{safe_audience}"
    docx_path = os.path.join(output_dir, f"{file_prefix}_{timestamp}.docx")
    
    doc = Document()
    doc.add_heading(module.get("topic", "Untitled Curriculum"), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Target Audience: {module.get('target_audience', 'N/A')}", style='Intense Quote')

    def add_markdown_content(heading, content_str, level):
        if not content_str or not isinstance(content_str, str): return
        doc.add_heading(heading, level=level)
        for line in content_str.strip().split('\n'):
            stripped_line = line.strip()
            if stripped_line.startswith('### '): doc.add_heading(stripped_line.replace('### ', ''), level=level + 2)
            elif stripped_line.startswith('## '): doc.add_heading(stripped_line.replace('## ', ''), level=level + 1)
            elif stripped_line.startswith('* '): doc.add_paragraph(stripped_line.replace('* ', '').strip(), style='List Bullet')
            elif stripped_line.startswith('- '): doc.add_paragraph(stripped_line.replace('- ', '').strip(), style='List Bullet')
            elif stripped_line:
                if '**' in stripped_line:
                    p = doc.add_paragraph()
                    for i, part in enumerate(stripped_line.split('**')):
                        p.add_run(part).bold = (i % 2 == 1)
                else:
                    doc.add_paragraph(stripped_line)

    add_markdown_content("Module Outline", module.get("module_outline"), 2)
    add_markdown_content("Learning Material", module.get("learning_material"), 2)
    add_markdown_content("Assessments", module.get("assessments"), 2)
    add_markdown_content("Resources", module.get("resources"), 2)
    
    doc.save(docx_path)
    print(f"📄 Saved DOCX to: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python main.py \"<Topic Name>\" \"<Target Audience>\"")
        sys.exit(1)
        
    topic = sys.argv[1]
    target_audience = sys.argv[2]
    
    print("\n🎓 Welcome to the Multi-Agent Curriculum Designer 🎓")
    print("-" * 60)
    
    try:
        designer = CurriculumDesignerSystem()
        final_module = designer.run_curriculum_design(topic, target_audience)
        
        if final_module and "raw_output" not in final_module:
            print("\n✅ Curriculum module created successfully!")
            save_to_json(final_module, topic, target_audience)
            save_to_docx(final_module, topic, target_audience)
            print("\n🎉 Process completed!")
        else:
            print("\n❌ Module creation failed or returned raw output.")
            print(json.dumps(final_module, indent=2))
            
    except Exception as e:
        print(f"\n❌ An unexpected error occurred: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)